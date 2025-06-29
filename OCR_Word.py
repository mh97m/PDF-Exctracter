from pdf2image import convert_from_path
import pytesseract
from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt
from docx.oxml import OxmlElement
import re

# Set Tesseract path and language
pytesseract.pytesseract.tesseract_cmd = '/usr/bin/tesseract'
lang = 'fas'  # Persian

def remove_parentheses_content(text):
    return re.sub(r'\(.*?\)', '', text)

def set_paragraph_rtl(paragraph):
    p = paragraph._element
    pPr = p.get_or_add_pPr()
    bidi = OxmlElement('w:bidi')
    bidi.set(qn('w:val'), '1')
    pPr.append(bidi)

def add_fonted_run(paragraph, text):
    """Add multiple runs to a paragraph, assigning different fonts based on script."""
    chunks = re.findall(r'[\u0600-\u06FF]+|[a-zA-Z0-9]+|[\s\.\,\:\;\!\?\-\_]+', text)
    for chunk in chunks:
        run = paragraph.add_run(chunk)
        run.font.size = Pt(14)

        # Force creating rPr and rFonts elements
        rPr = run._element.get_or_add_rPr()
        rFonts = rPr.find(qn('w:rFonts'))
        if rFonts is None:
            rFonts = OxmlElement('w:rFonts')
            rPr.append(rFonts)

        # Persian
        if re.search(r'[\u0600-\u06FF]', chunk):
            run.font.name = 'B Nazanin'
            rFonts.set(qn('w:ascii'), 'B Nazanin')
            rFonts.set(qn('w:hAnsi'), 'B Nazanin')
            rFonts.set(qn('w:eastAsia'), 'B Nazanin')
        else:
            run.font.name = 'Times New Roman'
            rFonts.set(qn('w:ascii'), 'Times New Roman')
            rFonts.set(qn('w:hAnsi'), 'Times New Roman')
            rFonts.set(qn('w:eastAsia'), 'Times New Roman')

def extract_persian_ocr_to_word(pdf_path, output_docx_path):
    images = convert_from_path(pdf_path)
    doc = Document()

    for i, image in enumerate(images, start=1):
        text = pytesseract.image_to_string(image, lang=lang)
        text = remove_parentheses_content(text).strip()

        title_paragraph = doc.add_paragraph(f"\n=== PAGE {i} ===")
        set_paragraph_rtl(title_paragraph)
        add_fonted_run(title_paragraph, "")

        if text:
            content_paragraph = doc.add_paragraph()
            set_paragraph_rtl(content_paragraph)
            add_fonted_run(content_paragraph, text)

    doc.save(output_docx_path)
    print(f"OCR Persian text saved to {output_docx_path}")

# Example usage
extract_persian_ocr_to_word(
    input('PDF file name (without .pdf): ') + ".pdf",
    input('Output DOCX file name (without .docx): ') + ".docx"
)
